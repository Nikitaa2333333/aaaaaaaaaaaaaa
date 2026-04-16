import docx
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("Буклет (2).docx")

lines = []
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style = para.style.name

    if style.startswith('Heading 1'):
        lines.append(f"\n# {text}")
    elif style.startswith('Heading 2'):
        lines.append(f"\n## {text}")
    elif style.startswith('Heading 3'):
        lines.append(f"\n### {text}")
    elif style.startswith('List'):
        lines.append(f"  - {text}")
    else:
        lines.append(text)

output = "\n".join(lines)
with open("content.txt", "w", encoding="utf-8") as f:
    f.write(output)

print(output)
print("\n\n--- Уникальные стили в документе ---")
styles_used = set()
for para in doc.paragraphs:
    if para.text.strip():
        styles_used.add(para.style.name)
for s in sorted(styles_used):
    print(s)
